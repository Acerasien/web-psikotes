import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Industrial Blue: #1E3A8A (30, 58, 138)
COLOR_BLUE_BG = "1E3A8A"
COLOR_BLUE_TEXT = RGBColor(21, 61, 99)
COLOR_GREY_BG = "F1F5F9"
COLOR_WHITE = "FFFFFF"

def set_cell_background(cell, fill_color):
    """Set background color for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    """Helper to set cell text with styling."""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = size
    if color:
        run.font.color.rgb = color
    return paragraph

def add_score_bar(cell, percentage, fill_color="3B82F6", height_pt=8, bg_color="E2E8F0"):
    """Adds a visual score bar inside a table cell."""
    # Only clear if the cell has one empty paragraph (default state)
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip():
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        
    # Create a 1x2 table inside the cell
    bar_table = cell.add_table(rows=1, cols=2)
    bar_table.autofit = False
    bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Calculate widths based on percentage
    total_width = cell.width if cell.width else Inches(1.5)
    usable_width = total_width * 0.95
    
    fill_width = usable_width * (max(0, min(100, percentage)) / 100)
    empty_width = max(1, usable_width - fill_width)
    
    # Set widths for columns and cells to force layout
    bar_table.columns[0].width = int(fill_width)
    bar_table.columns[1].width = int(empty_width)
    
    for row in bar_table.rows:
        row.cells[0].width = int(fill_width)
        row.cells[1].width = int(empty_width)
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_pt * 20)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        
        for c in row.cells:
            tcPr = c._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '0')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            
            for p in c.paragraphs:
                pPr = p._p.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), str(int(height_pt * 20)))
                spacing.set(qn('w:lineRule'), 'exact')
                pPr.append(spacing)
    
    set_cell_background(bar_table.cell(0, 0), fill_color)
    set_cell_background(bar_table.cell(0, 1), bg_color)

def add_section_heading(doc, text, size=Pt(11)):
    """Add a shaded section heading paragraph with a bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(14) # Reduced from 18
    p.space_after = Pt(4)  # Reduced from 6
    
    # Shading and Border
    pPr = p._p.get_or_add_pPr()
    
    # Shading: Dark Blue, Text 2, Lighter 90% (Approx #EEF2F7)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'C5D9F1')
    pPr.append(shd)
    
    # Bottom border
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6') # 0.75 pt
    bottom.set(qn('w:color'), '1F4E79') # Dark blue
    pbdr.append(bottom)
    pPr.append(pbdr)
    
    run = p.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = COLOR_BLUE_TEXT
    return p

# --- SHARED SUMMARY LOGIC ---
def get_summary_data(code, d):
    if not d: return ("-", "-", "-")
    if code == "DISC":
        g3 = d.get('graph_iii', {})
        dom = max(g3, key=g3.get) if g3 else "D"
        val = g3.get(dom, 0)
        dom_map = {"D":"Dominance","I":"Influence","S":"Steadiness","C":"Compliance"}
        return (f"{dom[0]} - {dom_map.get(dom, dom)} (G-III: {val})", "Tegas & Berorientasi Hasil", "SESUAI")
    if code == "TEMP":
        prim = d.get('primary', '-')
        # Mapping back to trait keys
        rev_map = {"Sanguine":"S", "Choleric":"C", "Melancholic":"M", "Phlegmatic":"P"}
        score = d.get('raw_scores', {}).get(rev_map.get(prim, ""), 0)
        return (f"{prim} Dominan ({score}/30)", "Analitis & Detail", "SESUAI")
    if code == "IQ":
        iq = d.get('iq', 0)
        cls = "Rata-rata"
        if iq >= 110: cls = "Di Atas Rata-rata"
        if iq >= 120: cls = "Unggul"
        if iq <= 90: cls = "Rata-rata Bawah"
        return (f"IQ {iq} - {cls}", "Kemampuan Kognitif Baik", "SESUAI")
    if code == "LOGIC":
        iq = d.get('est_iq', 0)
        cls = d.get('classification', 'Rata-rata')
        return (f"IQ {iq} - {cls}", "Logika Numerik Baik", "SESUAI")
    if code == "SPEED":
        score = d.get('score', 0)
        band = d.get('band', '-')
        # Mapping common band names to professional labels
        band_map = {"Excellent": "Tahan Tekanan", "Good": "Baik", "Average": "Cukup", "Needs Improvement": "Perlu Ditingkatkan"}
        clean_band = band_map.get(band, band)
        return (f"Skor {score} - {clean_band}", "Mampu dalam Kondisi Tekanan", "SESUAI")
    if code == "CBI":
        lvl_map = {"White": "Rendah", "Light Blue": "Sedang", "Dark Blue": "Tinggi", "-": "-"}
        lvl = lvl_map.get(d.get('overall_level', '-'), d.get('overall_level', '-'))
        return (f"Risiko {lvl} di Semua Aspek", "Minim Perilaku Kontraproduktif", "AMAN")
    if code == "LEAD":
        top_list = d.get('top_factors', [("N",0)])
        if not top_list: return ("-", "-", "-")
        top1_norm = top_list[0][0]
        val = d.get('stanines', {}).get(top1_norm, 0)
        # Short map for PAPI names
        papi_names = {"N":"Need to Lead","G":"Work Vigor","A":"Achievement","X":"Change","P":"Notice"}
        return (f"Top: {papi_names.get(top1_norm, top1_norm)} ({top1_norm}={val})", "Potensi Kepemimpinan Kuat", "UNGGUL")
    if code == "MEM":
        acc = d.get('accuracy', 0)
        score = d.get('correct_count', 0)
        total = d.get('total_answered', 0)
        band = "Sangat Baik" if acc >= 80 else "Baik" if acc >= 60 else "Rata-rata"
        if acc < 40: band = "Perlu Ditingkatkan"
        return (f"{score}/{total} ({acc}%) - {band}", "Kemampuan Mengingat Tinggi", "SESUAI")
    return ("-", "-", "-")

SUMMARY_CONFIGS = [
    ("DISC", "DISC"),
    ("TEMP", "4 Temperament"),
    ("IQ", "IQ Pola (CFIT)"),
    ("LOGIC", "IQ Aritmatika (WPT)"),
    ("SPEED", "Speed Test"),
    ("CBI", "CBI™"),
    ("LEAD", "Papikostick"),
    ("MEM", "Memory Test"),
]


def generate_participant_docx(user, results):
    """
    Generate Word report for a participant.
    user: User model instance
    results: list of Result model instances (latest for each test)
    """
    doc = Document()
    
    # --- GLOBAL STYLE (Abadi Font) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Abadi'
    font.size = Pt(10)
    
    # --- PAGE SETUP ---
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    
    # Margins
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    # --- REPEATING HEADER (all pages) ---
    header = section.header
    header.is_linked_to_previous = False
    
    # Remove all existing paragraphs in header
    for p in header.paragraphs:
        header.paragraphs[0]._element.getparent().remove(p._element)
    
    # Add new paragraph for the header image
    p_header = header.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = p_header.add_run()
    
    import os
    header_path = os.path.join(os.path.dirname(__file__), '..', '..', 'client', 'public', 'header.png')
    try:
        # Set width to 18.5 cm (full page width minus margins)
        run_header.add_picture(header_path, width=Cm(18.5))
    except Exception as e:
        # Fallback if image is missing
        run_header.add_text("ANDAMAS GROUP - PSYCHOLOGICAL ASSESSMENT REPORT")
        run_header.bold = True

    # --- REPEATING FOOTER (all pages) ---
    footer = section.footer
    footer.is_linked_to_previous = False
    p_foot = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("Confidential – Psychological Assessment Report | Andamas Group – HRGA Department | Page ")
    r_f.font.size = Pt(7)
    r_f.font.color.rgb = RGBColor(100, 100, 100)
    r_f.font.italic = True
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg = p_foot.add_run()
    run_pg.font.size = Pt(7)
    run_pg.font.color.rgb = RGBColor(100, 100, 100)
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)

    # --- DATA KARYAWAN ---
    completed_dates = [r.completed_at for r in results if r.completed_at]
    tgl_tes_str = min(completed_dates).strftime("%d/%m/%Y") if completed_dates else datetime.now().strftime("%d/%m/%Y")
    
    add_section_heading(doc, "DATA KARYAWAN", size=Pt(12))
    
    table_karyawan = doc.add_table(rows=4, cols=4)
    table_karyawan.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_karyawan.style = 'Table Grid'
    table_karyawan.autofit = False  # Disable autofit
    
    # Set exact widths
    widths = [Cm(2.99), Cm(6.0), Cm(3.75), Cm(5.75)]
    for row in table_karyawan.rows:
        row.height = Cm(0.75)
        for i, width in enumerate(widths):
            row.cells[i].width = width
    
    BG_LIGHT_BLUE = "DAE9F7"
    NAVY_TEXT = RGBColor(30, 77, 140)
    
    data = [
        ("Nama Lengkap", user.full_name or user.username, "Tanggal Tes", tgl_tes_str),
        ("Usia / Gender", f"{user.age or '-'} Tahun / {user.gender or '-'}", "Pendidikan Terakhir", user.education or "-"),
        ("Posisi Dilamar", user.position or "-", "Departement", user.department or "-"),
        ("Status", "Internal / Eksternal", "Assessor", "[Nama Assessor]")
    ]
    
    for i, row_data in enumerate(data):
        row = table_karyawan.rows[i]
        # Col 0: Label
        set_cell_background(row.cells[0], BG_LIGHT_BLUE)
        set_cell_text(row.cells[0], row_data[0], bold=True, color=NAVY_TEXT, size=Pt(7.5))
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Col 1: Value
        set_cell_text(row.cells[1], f": {row_data[1]}", size=Pt(9), color=RGBColor(0,0,0))
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Col 2: Label
        set_cell_background(row.cells[2], BG_LIGHT_BLUE)
        set_cell_text(row.cells[2], row_data[2], bold=True, color=NAVY_TEXT, size=Pt(7.5))
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Col 3: Value
        set_cell_text(row.cells[3], f": {row_data[3]}", size=Pt(9), color=RGBColor(0,0,0))
        row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    
    p_summary = doc.add_paragraph()
    p_summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_summary.space_after = Pt(12)
    run_sum = p_summary.add_run("Assessment Summary Result")
    run_sum.bold = True
    run_sum.font.size = Pt(14)
    run_sum.font.color.rgb = COLOR_BLUE_TEXT
    
    # Add "2 line thingy" (double border)
    pPr = p_summary._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double') # This is the "2 line thingy"
    bottom.set(qn('w:sz'), '12') # 1.5 pt
    bottom.set(qn('w:color'), '1F4E79')
    pbdr.append(bottom)
    pPr.append(pbdr)

    # Group results by code
    res_dict = {r.test.code: r for r in results if r.details}

    # --- 1. DISC ---
    if "DISC" in res_dict:
        add_section_heading(doc, "1 | DISC Personality Test")
        
        table_disc = doc.add_table(rows=5, cols=5)
        table_disc.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_disc.style = 'Table Grid'
        table_disc.autofit = False
        
        # Exact widths scaled to fit 18.5 cm usable width
        disc_widths = [Cm(5.70), Cm(3.20), Cm(3.20), Cm(3.20), Cm(3.20)]
        for r_idx, row in enumerate(table_disc.rows):
            if r_idx == 0: row.height = Cm(0.7)
            for i, width in enumerate(disc_widths):
                row.cells[i].width = width
        
        # Header heights
        table_disc.rows[0].height = Cm(0.7)
        table_disc.rows[1].height = Cm(0.9)
        
        # Header
        headers = ["Grafik DISC", "D", "I", "S", "C"]
        colors = ["156082", "C0392B", "E67E22", "27AE60", "2980B9"]
        for i, h in enumerate(headers):
            cell = table_disc.cell(0, i)
            set_cell_background(cell, colors[i])
            set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        # Subheader
        sub_first = table_disc.cell(1, 0)
        set_cell_text(sub_first, "Dimensi (Dominance / Influence / Steadiness / Compliance)", size=Pt(7), align=WD_ALIGN_PARAGRAPH.LEFT).runs[0].font.italic = True
        sub_first.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        subs_trait = ["Dominance", "Influence", "Steadiness", "Compliance"]
        for i, s in enumerate(subs_trait):
            cell = table_disc.cell(1, i+1)
            set_cell_text(cell, s, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.italic = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        disc_details = res_dict["DISC"].details
        
        # Row 2,3,4: Most, Least, Combined
        rows_data = [
            ("Grafik I – MOST", disc_details.get("graph_i", {})),
            ("Grafik II – LEAST", disc_details.get("graph_ii", {})),
            ("Grafik III – COMBINED", disc_details.get("graph_iii", {}))
        ]        
        BG_HIGHLIGHT = "FFFBEB" # Light yellow for high score
        
        for r_idx, (label, data_map) in enumerate(rows_data):
            row = table_disc.rows[r_idx + 2]
            row.height = Cm(1.0) # Exact height from spec
            
            set_cell_background(row.cells[0], COLOR_GREY_BG)
            set_cell_text(row.cells[0], label, bold=True, color=COLOR_BLUE_TEXT, size=Pt(9))
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Find max trait for highlighting
            traits = ["D", "I", "S", "C"]
            max_trait = max(traits, key=lambda t: abs(data_map.get(t, 0)))
            
            bar_colors = {"D": "DC2626", "I": "F59E0B", "S": "10B981", "C": "3B82F6"}
            bg_colors = {"D": "FEE2E2", "I": "FEF3C7", "S": "D1FAE5", "C": "DBEAFE"}
            
            for c_idx, trait in enumerate(traits):
                val = data_map.get(trait, 0)
                cell = row.cells[c_idx + 1]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Highlight if max
                if trait == max_trait:
                    set_cell_background(cell, BG_HIGHLIGHT)
                
                # Set text color to match trait
                t_color = RGBColor.from_string(bar_colors[trait])
                p_score = set_cell_text(cell, str(val), bold=True, color=t_color, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11))
                # Add 2pt space after score to separate from bar
                p_score.paragraph_format.space_after = Pt(2)
                
                # Specialized bar with tinted background
                add_score_bar(cell, (abs(val) / 24 * 100), fill_color=bar_colors[trait], bg_color=bg_colors[trait])

        # Kesimpulan / Tipe Dominan Box
        row_kes = table_disc.add_row()
        row_kes.height = Cm(0.34) # Exact height from spec
        cell_kes = row_kes.cells[0]
        cell_kes.merge(table_disc.rows[-1].cells[-1])
        set_cell_background(cell_kes, "FEF3C7") # Summary box color
        cell_kes.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        g3 = disc_details.get("graph_iii", {})
        primary_trait = max(g3, key=g3.get) if g3 else "D"
        trait_names = {"D": "Dominance", "I": "Influence", "S": "Steadiness", "C": "Compliance"}
        trait_desc = {
            "D": "Percaya diri, menyukai tantangan, berorientasi pada pencapaian, senang memberi pengaruh besar dalam hasil.",
            "I": "Hangat, optimis, terbuka, mudah membangun jaringan sosial, dan termotivasi lewat interaksi.",
            "S": "Sabar, konsisten, setia, menghargai kestabilan, dan nyaman di lingkungan harmonis.",
            "C": "Teliti, berhati-hati, mengutamakan akurasi, nyaman dengan aturan jelas dan standar tinggi."
        }
        
        # Dictionaries for interpretations
        desc_i = {
            "D": "Cenderung berorientasi pada hasil dan pencapaian target. Mereka tegas, cepat dalam mengambil keputusan, menyukai tantangan, dan senang memimpin. Biasanya berperan mendorong tim agar tetap bergerak maju dan produktif.",
            "I": "Energik, persuasif, dan mudah membangun hubungan. Mereka menyukai interaksi sosial, mampu memotivasi orang lain, serta menciptakan suasana kerja yang positif dan penuh semangat.",
            "S": "Stabil, sabar, dan pendengar yang baik. Mereka berfokus pada kerja sama tim, menjaga keharmonisan, serta lebih nyaman bekerja di lingkungan yang teratur dan konsisten.",
            "C": "Teliti, logis, dan analitis. Mereka bekerja dengan standar tinggi, mengikuti aturan, serta fokus pada detail dan kualitas hasil."
        }
        desc_ii = {
            "D": "Cenderung menjadi lebih keras, tidak sabar, bahkan mendominasi. Mereka ingin segera menyelesaikan masalah dengan cepat, meskipun kadang mengabaikan detail atau perasaan orang lain (Dominance - D).",
            "I": "Bisa menjadi terlalu emosional, kurang fokus, dan mudah terdistraksi. Mereka sering mencari dukungan dari orang lain untuk merasa aman (Influence - I).",
            "S": "Menjadi lebih pasif, menghindari konflik, dan membutuhkan arahan yang jelas. Mereka bisa merasa kewalahan jika perubahan terlalu cepat (Steadiness - S).",
            "C": "Cenderung perfeksionis, terlalu banyak menganalisis, dan takut membuat kesalahan. Hal ini dapat memperlambat pengambilan keputusan (Conscientiousness - C)."
        }
        desc_iii = {
            "D": "Secara natural, mereka percaya diri, berorientasi pada pencapaian, dan menyukai tantangan baru. Mereka senang berada dalam posisi yang memberi pengaruh besar terhadap hasil.",
            "I": "Alami sebagai pribadi yang hangat, optimis, terbuka, dan mudah membangun jaringan sosial. Mereka lebih termotivasi ketika bisa berinteraksi dengan orang lain.",
            "S": "Sifat alaminya adalah sabar, konsisten, dan setia. Mereka menghargai kestabilan, serta lebih suka bekerja dalam lingkungan yang harmonis dan minim konflik.",
            "C": "Alami sebagai individu yang teliti, berhati-hati, dan mengutamakan akurasi. Mereka nyaman bekerja dengan aturan yang jelas serta standar kerja yang tinggi."
        }

        g1 = disc_details.get("graph_i", {})
        g2 = disc_details.get("graph_ii", {})
        g3 = disc_details.get("graph_iii", {})
        
        p1 = max(g1, key=g1.get) if g1 else "D"
        p2 = max(g2, key=g2.get) if g2 else "D"
        p3 = max(g3, key=g3.get) if g3 else "D"

        kesimpulan_text = f"Kesimpulan: {desc_i.get(p1, '')} {desc_ii.get(p2, '')} {desc_iii.get(p3, '')}"

        # 1st Paragraph: Kesimpulan (Dynamic)
        p_kes1 = cell_kes.paragraphs[0]
        p_kes1.paragraph_format.left_indent = Pt(8)
        p_kes1.paragraph_format.space_before = Pt(6) # Padding top
        r_kes = p_kes1.add_run(kesimpulan_text)
        r_kes.font.size = Pt(8)
        r_kes.font.italic = True
        
        # 2nd Paragraph: Tipe Dominan Aktual
        p_kes2 = cell_kes.add_paragraph()
        p_kes2.paragraph_format.left_indent = Pt(8)
        p_kes2.paragraph_format.space_after = Pt(6) # Padding bottom
        
        r1 = p_kes2.add_run(f"Tipe Dominan Aktual (G-III): ")
        r1.bold = True
        r1.font.size = Pt(8)
        r1.font.color.rgb = COLOR_BLUE_TEXT
        
        r_trait = p_kes2.add_run(f"[{primary_trait} – {trait_names.get(primary_trait, '')}]")
        r_trait.bold = True
        r_trait.font.size = Pt(8)
        r_trait.font.color.rgb = RGBColor.from_string(bar_colors.get(primary_trait, "DC2626"))
        
        p_kes2.add_run(f" → {trait_desc.get(primary_trait, '')}").font.size = Pt(8)
        
        # Gap after DISC (Safe here as TEMP follows)
        doc.add_paragraph()

    # --- 2. TEMPERAMENT ---
    if "TEMP" in res_dict:
        add_section_heading(doc, "2 | Four Temperament Test")
        
        table_temp = doc.add_table(rows=5, cols=4)
        table_temp.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_temp.style = 'Table Grid'
        table_temp.autofit = False
        
        # Exact widths and height for header
        temp_widths = [Cm(4.16), Cm(2.30), Cm(3.04), Cm(9.04)]
        for r_idx, row in enumerate(table_temp.rows):
            if r_idx == 0: row.height = Cm(0.75)
            for i, width in enumerate(temp_widths):
                row.cells[i].width = width
        
        headers = ["Tipe Temperamen", "Skor", "Grafik", "Karakteristik Utama"]
        for i, h in enumerate(headers):
            cell = table_temp.cell(0, i)
            set_cell_background(cell, "156082")
            set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        temp_details = res_dict["TEMP"].details
        scores = temp_details.get("raw_scores", {})
        
        traits = [
            ("Sanguine", "S", "Ekspresif, antusias, mudah bergaul, spontan"),
            ("Melancholic", "M", "Analitis, terstruktur, perfeksionis, detail-oriented"),
            ("Choleric", "C", "Tegas, pemimpin, goal-oriented, cepat bertindak"),
            ("Phlegmatic", "P", "Tenang, sabar, stabil, diplomatis")
        ]
        
        for r_idx, (name, key, desc) in enumerate(traits):
            row = table_temp.rows[r_idx + 1]
            row.height = Cm(0.7) # Added height consistency
            is_primary = (name == temp_details.get("primary"))
            val = scores.get(key, 0)
            
            # Name
            c0 = row.cells[0]
            set_cell_text(c0, f"{name}{' ★ Dominan' if is_primary else ''}", bold=is_primary, color=COLOR_BLUE_TEXT if is_primary else None, size=Pt(9))
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Score
            c1 = row.cells[1]
            set_cell_text(c1, str(val), bold=True, color=RGBColor(220,38,38) if is_primary else None, align=WD_ALIGN_PARAGRAPH.CENTER)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Bar
            c2 = row.cells[2]
            temp_colors = {"Sanguine": "F59E0B", "Melancholic": "3B82F6", "Choleric": "DC2626", "Phlegmatic": "10B981"}
            bg_colors = {"Sanguine": "FEF3C7", "Melancholic": "DBEAFE", "Choleric": "FEE2E2", "Phlegmatic": "D1FAE5"}
            add_score_bar(c2, (val / 30) * 100, fill_color=temp_colors.get(name, COLOR_BLUE_BG), bg_color=bg_colors.get(name, "E2E8F0"))
            c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Desc
            c3 = row.cells[3]
            set_cell_text(c3, desc, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT).runs[0].font.italic = True
            c3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        # Gap after TEMP (Safe here as IQ follows)
        doc.add_paragraph()


    # --- 3 & 4. IQ / LOGIC ---
    if "IQ" in res_dict or "LOGIC" in res_dict:
        add_section_heading(doc, "3 & 4 | Tes Kemampuan Intelektual (IQ Pola – CFIT & IQ Aritmatika dan Logika – WPT)")
        
        table_iq = doc.add_table(rows=3, cols=6)
        table_iq.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_iq.style = 'Table Grid'
        table_iq.autofit = False
        
        # Row Heights from spec
        table_iq.rows[0].height = Cm(0.7) # Header
        table_iq.rows[1].height = Cm(1.1) # CFIT
        table_iq.rows[2].height = Cm(1.2) # WPT
        
        # Exact widths from spec
        iq_widths = [Cm(1.74), Cm(3.50), Cm(1.75), Cm(1.25), Cm(2.75), Cm(7.50)]
        for row in table_iq.rows:
            for i, width in enumerate(iq_widths):
                row.cells[i].width = width
        
        headers = ["Kode", "Nama Tes", "Skor Mentah", "Skor IQ", "Kategori", "Interpretasi"]
        for i, h in enumerate(headers):
            cell = table_iq.cell(0, i)
            set_cell_background(cell, "156082")
            set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        # CFIT Row
        cfit_row = table_iq.rows[1]
        # set_cell_background(cfit_row.cells[0], "EAF2FB")
        cfit_row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_text(cfit_row.cells[0], "CFIT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Nama Tes: Title (Center) + Detail (Left)
        c1 = cfit_row.cells[1]
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # set_cell_background(cfit_row.cells[1], "EAF2FB")
        p_title = c1.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p_title.add_run("Culture Fair Intelligence Test")
        r1.bold = True
        r1.font.size = Pt(8)
        
        p_detail = c1.add_paragraph()
        p_detail.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_detail.paragraph_format.left_indent = Pt(5)
        r2 = p_detail.add_run("Test IQ Pola / Non-Verbal")
        r2.font.size = Pt(7)
        r2.font.italic = True
        
        if "IQ" in res_dict:
            res_iq = res_dict["IQ"]
            det = res_iq.details or {}
            # Total raw score across all phases
            raw_score = det.get("raw_score", res_iq.score)
            
            # Logic from scoring/iq.py
            scaled = round(raw_score / 2)
            from scoring.iq import SCORE_TO_IQ, classify_iq
            iq_val = SCORE_TO_IQ.get(scaled, 0)
            category = classify_iq(iq_val)

            # Skor Mentah (Total Phase 1-8)
            set_cell_text(cfit_row.cells[2], str(raw_score), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            cfit_row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # set_cell_background(cfit_row.cells[2], "EAF2FB")
            
            # Skor IQ
            set_cell_text(cfit_row.cells[3], str(iq_val), bold=True, color=COLOR_BLUE_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
            cfit_row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # set_cell_background(cfit_row.cells[3], "EAF2FB")
            
            # Kategori
            set_cell_text(cfit_row.cells[4], category, bold=True, color=COLOR_BLUE_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
            cfit_row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # set_cell_background(cfit_row.cells[4], "EAF2FB") 
            
            # Keterangan (Left)
            set_cell_text(cfit_row.cells[5], "Kemampuan penalaran abstrak, mengenal pola, dan logika non-verbal.", size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).runs[0].font.italic = True
            cfit_row.cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # set_cell_background(cfit_row.cells[5], "EAF2FB")

        # WPT Row
        wpt_row = table_iq.rows[2]
        wpt_row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_text(wpt_row.cells[0], "WPT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Nama Tes: Title (Center) + Detail (Left)
        c1w = wpt_row.cells[1]
        c1w.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_title2 = c1w.paragraphs[0]
        p_title2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1w = p_title2.add_run("Wonderlic Personnel Test")
        r1w.bold = True
        r1w.font.size = Pt(8)
        
        p_detail2 = c1w.add_paragraph()
        p_detail2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_detail2.paragraph_format.left_indent = Pt(5)
        r2w = p_detail2.add_run("IQ Aritmatika di Logika Verbal")
        r2w.font.size = Pt(7)
        r2w.font.italic = True
        
        if "LOGIC" in res_dict:
            det = res_dict["LOGIC"].details
            # Skor Mentah
            set_cell_text(wpt_row.cells[2], str(det.get("correct_count", 0)), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            # Skor IQ
            set_cell_text(wpt_row.cells[3], str(det.get("est_iq", "-")), bold=True, color=COLOR_BLUE_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
            # Kategori
            set_cell_text(wpt_row.cells[4], str(det.get("classification", "-")), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            # Keterangan (Left)
            set_cell_text(wpt_row.cells[5], "Kemampuan numerik, penalaran verbal, dan pemecahan masalah dalam konteks kerja.", size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).runs[0].font.italic = True
            
        for i in [2,3,4,5]: wpt_row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # IQ Reference Scale - Styled as a small caption to avoid overflow
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(0)
        p_ref.paragraph_format.space_after = Pt(0)
        p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_ref = p_ref.add_run("Skala IQ Referensi: < 80 = Rendah  |  80–89 = Di Bawah Rata-Rata  |  90–109 = Rata-Rata  |  110–119 = Di Atas Rata-Rata  |  ≥ 120 = Superior")
        r_ref.font.size = Pt(7)
        r_ref.bold = True
        r_ref.font.color.rgb = RGBColor(100, 100, 100)
        
        # Keep it attached to the table
        p_ref.paragraph_format.keep_with_next = False 
        # Actually, let's make sure the WPT row keeps with next
        wpt_row.cells[0].paragraphs[0].paragraph_format.keep_with_next = True

    # --- PAGE BREAK ---
    doc.add_page_break()

    # --- 5. SPEED ---
    if "SPEED" in res_dict:
        add_section_heading(doc, "5 | Speed Tes – Ketahanan terhadap tekanan")
        
        # 3 Rows, 6 Columns
        table_speed = doc.add_table(rows=3, cols=6)
        table_speed.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_speed.style = 'Table Grid'
        table_speed.autofit = False
        
        # Exact widths from spec
        speed_widths = [Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.2), Cm(7.8)]
        for r_idx, row in enumerate(table_speed.rows):
            if r_idx == 1: row.height = Cm(1.5)  # Category labels height
            if r_idx == 2: row.height = Cm(0.85) # Merged score bar height
            for i, width in enumerate(speed_widths):
                row.cells[i].width = width
        
        ranges = ["0 - 30", "31-60", "61- 80", "81- 90", "91 - 100"]
        colors = ["DC2626", "D97706", "10B981", "059669", "1E3A8A"]
        
        # Row 0: Headers
        table_speed.rows[0].height = Cm(0.5)
        for i, h in enumerate(ranges):
            cell = table_speed.cell(0, i)
            set_cell_background(cell, colors[i])
            set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Header for Interpretation
        cell_h_int = table_speed.cell(0, 5)
        set_cell_background(cell_h_int, "156082")
        set_cell_text(cell_h_int, "Catatan Interpretasi", bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
        cell_h_int.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        labels = [
            "Kurang Siap Menghadapi Tekanan", 
            "Cukup Siap Menghadapi Tekanan", 
            "Siap Menghadapi Tekanan", 
            "Sangat Siap Menghadapi Tekanan", 
            "Unggul dalam Tekanan Tinggi"
        ]
        
        det = res_dict["SPEED"].details
        score = det.get("score", 0)
        
        # Determine index
        idx = 0
        if score > 30: idx = 1
        if score > 60: idx = 2
        if score > 80: idx = 3
        if score > 90: idx = 4
        
        # Row 1: Category Labels
        for i, l in enumerate(labels):
            cell = table_speed.cell(1, i)
            is_active = (i == idx)
            # Tinted background for active category
            if is_active:
                bg_tint = ["FEE2E2", "FEF3C7", "D1FAE5", "D1FAE5", "DBEAFE"][i]
                set_cell_background(cell, bg_tint)
            
            set_cell_text(cell, l, bold=is_active, color=RGBColor.from_string(colors[i]) if is_active else RGBColor(120,120,120), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Interpretation logic
        interpretations = [
            "Individu memiliki ketahanan yang rendah terhadap tekanan kerja yang intens. Membutuhkan waktu lebih lama untuk beradaptasi dengan beban kerja tinggi.",
            "Individu cukup mampu bekerja dalam kondisi tekanan standar, namun performa mungkin menurun saat beban kerja meningkat drastis.",
            "Individu mampu bekerja dalam kondisi tekanan waktu dan beban kerja tinggi. Cocok untuk lingkungan operasional pertambangan yang dinamis.",
            "Individu memiliki ketahanan yang sangat baik terhadap tekanan. Tetap tenang dan produktif meskipun dalam situasi kritis.",
            "Individu memiliki kemampuan luar biasa dalam mengelola tekanan. Sangat efektif dalam pengambilan keputusan cepat di bawah situasi ekstrem."
        ]
        
        # Merged Interpretation Cell (Row 1-2, Col 5)
        cell_int = table_speed.cell(1, 5)
        cell_int.merge(table_speed.cell(2, 5))
        set_cell_text(cell_int, interpretations[idx], size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT).runs[0].font.italic = True
        cell_int.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Row 2: Score Bar (Merged 0-4)
        cell_score = table_speed.cell(2, 0)
        for col in range(1, 5):
            cell_score.merge(table_speed.cell(2, col))
        
        set_cell_background(cell_score, "F1F5F9")
        set_cell_text(cell_score, f"◀ SKOR: {score}", bold=True, color=RGBColor.from_string(colors[idx]), align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_score.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()

    # --- 6. CBI ---
    if "CBI" in res_dict:
        add_section_heading(doc, "6 | CBI™ – Counterproductive Behavior Indeks")
        
        det = res_dict["CBI"].details
        concerns = det.get("primary_concerns", {})
        
        table_cbi = doc.add_table(rows=len(concerns)+2, cols=4)
        table_cbi.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_cbi.style = 'Table Grid'
        table_cbi.autofit = False
        
        # Exact widths from spec
        cbi_widths = [Cm(3.74), Cm(2.25), Cm(2.25), Cm(10.26)]
        for row in table_cbi.rows:
            for i, width in enumerate(cbi_widths):
                row.cells[i].width = width
        
        headers = ["Aspek CBI™", "Skor", "Tingkat Risiko", "Keterangan"]
        table_cbi.rows[0].height = Cm(0.5)
        for i, h in enumerate(headers):
            cell = table_cbi.cell(0, i)
            set_cell_background(cell, "156082")
            set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        # CBI Config: Max scores and Professional descriptions
        cbi_config = {
            "Dependability": {
                "max": 22,
                "Sedang": "Perlu pemantauan terhadap tingkat kehadiran dan ketepatan waktu kerja.",
                "Tinggi": "Risiko tinggi terhadap perilaku mangkir atau ketidakhadiran tanpa izin.",
                "Rendah": "Risiko rendah terhadap ketidakhadiran; menunjukkan kedisiplinan yang baik."
            },
            "Aggression": {
                "max": 24,
                "Sedang": "Terdapat indikasi reaktivitas emosional yang perlu dikelola lebih lanjut.",
                "Tinggi": "Risiko tinggi terhadap perilaku impulsif atau konfrontasi di tempat kerja.",
                "Rendah": "Mampu mengelola emosi dengan baik dalam situasi kerja yang menantang."
            },
            "Substance Abuse": {
                "max": 22,
                "Sedang": "Perlu kewaspadaan terhadap potensi pengaruh zat eksternal pada produktivitas.",
                "Tinggi": "Indikasi kuat adanya risiko penyalahgunaan zat yang membahayakan keselamatan.",
                "Rendah": "Tidak ditemukan indikasi penggunaan zat terlarang yang memengaruhi kinerja."
            },
            "Honesty": {
                "max": 23,
                "Sedang": "Terdapat celah dalam pemahaman integritas yang membutuhkan pengawasan.",
                "Tinggi": "Risiko tinggi terhadap tindakan tidak jujur atau pelanggaran etika kerja.",
                "Rendah": "Menunjukkan integritas yang baik dan kepatuhan terhadap nilai-nilai kejujuran."
            },
            "Computer Abuse": {
                "max": 17,
                "Sedang": "Terdapat kecenderungan penggunaan aset digital untuk keperluan non-pekerjaan.",
                "Tinggi": "Risiko penyalahgunaan aset digital yang dapat merugikan keamanan data.",
                "Rendah": "Menggunakan fasilitas teknologi perusahaan secara bijak dan profesional."
            },
            "Sexual Harassment": {
                "max": 22,
                "Sedang": "Perlu edukasi mengenai batasan perilaku yang pantas di lingkungan kerja.",
                "Tinggi": "Risiko tinggi terhadap perilaku pelecehan atau pelanggaran batas profesional.",
                "Rendah": "Menjaga batas profesional dan menghargai etika interaksi antar rekan kerja."
            }
        }
        
        row_idx = 1
        for trait, c_data in concerns.items():
            if trait not in cbi_config: continue
            row = table_cbi.rows[row_idx]
            row.height = Cm(0.85)
            
            c_score = c_data.get('score', 0)
            c_level = c_data.get('level', 'White')
            max_val = cbi_config[trait]["max"]
            
            # Map level to Indo
            level_map = {
                "White": ("Rendah", RGBColor(16,185,129)),
                "Light Blue": ("Sedang", RGBColor(217,119,6)),
                "Dark Blue": ("Tinggi", RGBColor(220,38,38))
            }
            lvl_str, color = level_map.get(c_level, ("Rendah", RGBColor(16,185,129)))

            # Col 0: Aspek (Left)
            set_cell_text(row.cells[0], trait, bold=True, color=COLOR_BLUE_TEXT, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
            
            # Col 1: Score (Center) + Small bar matched to level color
            c1 = row.cells[1]
            set_cell_text(c1, f"{c_score}/{max_val}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Map level to Hex for bar
            bar_hex_map = {"White": "10B981", "Light Blue": "D97706", "Dark Blue": "DC2626"}
            bar_color = bar_hex_map.get(c_level, "10B981")
            add_score_bar(c1, (c_score / max_val) * 100, fill_color=bar_color, height_pt=5)
            
            # Col 2: Level (Center)
            set_cell_text(row.cells[2], lvl_str, bold=True, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Col 3: Keterangan (Left)
            desc_text = cbi_config[trait].get(lvl_str, "")
            set_cell_text(row.cells[3], desc_text, size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
            row.cells[3].paragraphs[0].runs[0].font.italic = True
            
            for cell in row.cells: cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_idx += 1
            
        # Catatan Row
        row_cat = table_cbi.rows[row_idx]
        row_cat.height = Cm(0.5)
        cell_cat = row_cat.cells[0]
        cell_cat.merge(row_cat.cells[-1])
        set_cell_background(cell_cat, "FEF3C7")
        set_cell_text(cell_cat, "Catatan: Skor CBI adalah skor risiko = semakin RENDAH skor, semakin BAIK hasilnya.", size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
        cell_cat.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()

    # --- 7. PAPI ---
    if "LEAD" in res_dict:
        add_section_heading(doc, "7 | PAPIKOSTICK – Top 5 Faktor Kepemimpinan")
        
        det = res_dict["LEAD"].details
        percentages = det.get("percentages", {})
        descriptions = det.get("norm_descriptions", {})
        
        # Sort and get top 5
        top5 = sorted(percentages.items(), key=lambda x: x[1], reverse=True)[:5]
        
        table_papi = doc.add_table(rows=len(top5)+1, cols=6)
        table_papi.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_papi.style = 'Table Grid'
        table_papi.autofit = False
        
        # Scaled widths to fit A4 (Sum ~18.5cm)
        # Ratio based on: Kode:1.4, Aspek:3.7, Score:3, Kategori:2.75, Deskripsi:7.5
        papi_widths = [Cm(0.63), Cm(1.36), Cm(3.60), Cm(2.92), Cm(2.68), Cm(7.31)]
        for r_idx, row in enumerate(table_papi.rows):
            if r_idx == 0: row.height = Cm(0.6) # Header height
            for i, width in enumerate(papi_widths):
                row.cells[i].width = width
        
        headers = ["#", "Kode", "Aspek Utama", "Score", "Kategori", "Deskripsi Interpretasi"]
        for i, h in enumerate(headers):
            cell = table_papi.cell(0, i)
            set_cell_background(cell, "156082")
            # Align: Center for all except Deskripsi Interpretasi (idx 5) which is Center Left
            align = WD_ALIGN_PARAGRAPH.CENTER
            if i == 5: align = WD_ALIGN_PARAGRAPH.LEFT
            
            p = set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=align, size=Pt(7.5))
            if i == 5: p.paragraph_format.left_indent = Pt(5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        theme_colors = ["0F172A", "2563EB", "10B981", "D97706", "7C3AED"]
        
        # Professional descriptions mapping
        papi_descs = {
            "G": "Memiliki energi dan ketahanan kerja yang tinggi, mampu menjaga performa optimal dalam jangka panjang.",
            "L": "Memiliki keinginan kuat untuk memimpin, mengorganisir, dan memberikan arahan kepada orang lain secara efektif.",
            "I": "Menunjukkan kemampuan pengambilan keputusan yang cepat dan berani mengambil inisiatif dalam tugas.",
            "A": "Dorongan kuat untuk mencapai target tinggi dan melampaui standar performa yang telah ditetapkan.",
            "N": "Memiliki keinginan untuk menyelesaikan tugas secara mandiri dan tuntas dengan pengawasan minimal.",
            "S": "Menunjukkan kebutuhan untuk menjalin interaksi sosial dan membangun hubungan kerja yang harmonis.",
            "X": "Terbuka terhadap perubahan, fleksibel dalam menghadapi situasi baru dan adaptif terhadap tantangan.",
            "P": "Memiliki motivasi untuk diakui dan mendapatkan apresiasi atas kontribusi profesional yang diberikan.",
            "R": "Memiliki kecenderungan untuk mematuhi aturan dan mengikuti prosedur kerja secara disiplin.",
            "D": "Menunjukkan ketelitian tinggi dalam bekerja dan sangat memperhatikan detail operasional.",
            "O": "Kebutuhan tinggi akan keteraturan dan kerapian dalam lingkungan maupun proses kerja.",
            "C": "Menunjukkan sikap pengendalian diri yang baik dan mampu bersikap tenang di bawah tekanan.",
            "E": "Mampu mengekspresikan emosi secara terbuka namun tetap dalam batas profesionalisme kerja.",
            "B": "Menunjukkan kebutuhan untuk diterima dalam kelompok dan menjadi bagian dari tim yang solid.",
            "F": "Memiliki sikap loyal dan suportif terhadap pimpinan serta otoritas di lingkungan kerja.",
            "W": "Kebutuhan tinggi akan arahan dan bimbingan yang jelas untuk meminimalisir kesalahan kerja.",
            "K": "Memiliki kecenderungan untuk mempertahankan argumen atau bersikap asertif dalam diskusi.",
            "Z": "Menunjukkan keinginan untuk terus bergerak dan dinamis dalam menyelesaikan tanggung jawab.",
            "T": "Memiliki ritme kerja yang cepat dan reaktif terhadap permintaan tugas yang mendesak.",
            "V": "Memiliki ketertarikan untuk bekerja secara dinamis dan aktif secara fisik maupun mental."
        }
        
        for i, (norm, pct) in enumerate(top5):
            row = table_papi.rows[i+1]
            row.height = Cm(0.85)
            t_color = theme_colors[i % len(theme_colors)]
            
            # Col 0: # (Rank Background)
            set_cell_background(row.cells[0], t_color)
            set_cell_text(row.cells[0], str(i+1), bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Col 1: Kode (Theme Color Text, Center)
            set_cell_text(row.cells[1], norm, bold=True, color=RGBColor.from_string(t_color), align=WD_ALIGN_PARAGRAPH.CENTER)
            
            desc_full = descriptions.get(norm, "")
            # Split "Name — Indonesian desc" format
            if "—" in desc_full:
                parts = desc_full.split("—")
                name = parts[0].strip()
            else:
                name = desc_full
            
            # Col 2: Aspek Utama (Center Left)
            set_cell_text(row.cells[2], name, bold=True, color=COLOR_BLUE_TEXT, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
            
            # Use stanine for x/9 score
            stanine = det.get("stanines", {}).get(norm, 5)
            # Col 3: Score (Theme Color, Center) + Bar
            c3 = row.cells[3]
            set_cell_text(c3, f"{stanine}/9", bold=True, color=RGBColor.from_string(t_color), align=WD_ALIGN_PARAGRAPH.CENTER)
            add_score_bar(c3, (stanine / 9) * 100, fill_color=t_color, height_pt=5)
            
            # Col 4: Kategori (Center)
            kategori = "Rata-rata"
            if stanine >= 7: kategori = "Tinggi"
            if stanine <= 3: kategori = "Rendah"
            set_cell_text(row.cells[4], kategori, align=WD_ALIGN_PARAGRAPH.CENTER)

            # Col 5: Deskripsi (Center Left)
            short_desc = papi_descs.get(norm, desc_full)
            set_cell_text(row.cells[5], short_desc, size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
            row.cells[5].paragraphs[0].runs[0].font.italic = True
            
            for cell in row.cells: cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()

    # --- 8. MEMORY ---
    if "MEM" in res_dict:
        add_section_heading(doc, "8 | Memory Test – Kemampuan Mengingat")
        
        table_mem = doc.add_table(rows=2, cols=4)
        table_mem.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_mem.style = 'Table Grid'
        table_mem.autofit = False
        
        # Exact widths from spec
        mem_widths = [Cm(4.35), Cm(2.64), Cm(2.50), Cm(9.03)]
        for r_idx, row in enumerate(table_mem.rows):
            if r_idx == 0: row.height = Cm(0.75) # Header height
            for i, width in enumerate(mem_widths):
                row.cells[i].width = width
        
        headers = ["Aspek", "Skor", "Kategori", "Interpretasi"]
        header_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for i, h in enumerate(headers):
            cell = table_mem.cell(0, i)
            set_cell_background(cell, "156082")
            set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=header_aligns[i], size=Pt(7.5))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if header_aligns[i] == WD_ALIGN_PARAGRAPH.LEFT:
                cell.paragraphs[0].paragraph_format.left_indent = Pt(5)
            
        row = table_mem.rows[1]
        row.height = Cm(2.0) # Detail row height
        
        set_cell_background(row.cells[0], COLOR_GREY_BG)
        set_cell_text(row.cells[0], "Kemampuan Memori (Short & Working Memory)", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
        
        det = res_dict["MEM"].details
        score = det.get("correct_count", 0)
        total = det.get('total_answered', 0) or 50 # Fallback if missing
        pct = det.get("accuracy", 0)
        band = det.get("band", "-")
        
        # Color mapping based on performance
        band_config = {
            "Excellent Memory": ("Sangat Baik", "10B981", "Memiliki kapasitas memori sangat kuat dalam mengolah dan mengingat detail informasi baru secara akurat."),
            "Good": ("Baik", "34D399", "Kapasitas memori baik; mampu mengolah dan menyimpan informasi baru dengan efektivitas yang terjaga."),
            "Average": ("Rata-rata", "D97706", "Kemampuan memori dalam batas normal; cukup mampu mengolah informasi standar di lingkungan kerja."),
            "Needs Improvement": ("Perlu Ditingkatkan", "DC2626", "Membutuhkan perhatian lebih dalam mengolah informasi; kapasitas memori perlu dilatih lebih lanjut.")
        }
        
        lvl_name, lvl_color, lvl_desc = band_config.get(band, (band, "10B981", "Kemampuan mengingat informasi."))
        
        # Col 1: Score (Center) + Bar
        c1 = row.cells[1]
        c1.paragraphs[0].clear()
        
        # 1. Raw Score (Large & Bold)
        p1 = c1.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(2)
        run1 = p1.add_run(f"{score}/{total}")
        run1.font.bold = True
        run1.font.size = Pt(12)
        run1.font.color.rgb = COLOR_BLUE_TEXT
        
        # 2. Percentage (Smaller)
        p2 = c1.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8) # Gap before bar
        run2 = p2.add_run(f"({pct}%)")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(100, 100, 100)
        
        # 3. Progress Bar (with padding)
        add_score_bar(c1, pct, fill_color=lvl_color, height_pt=7)
        
        # Col 2: Kategori (Center)
        set_cell_text(row.cells[2], lvl_name, bold=True, color=RGBColor.from_string(lvl_color), align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 3: Interpretasi (Left)
        set_cell_text(row.cells[3], lvl_desc, size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
        row.cells[3].paragraphs[0].runs[0].font.italic = True
        
        for cell in row.cells: cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- PAGE BREAK ---
    doc.add_page_break()

    # --- SUMMARY PAGE ---
    add_section_heading(doc, "RINGKASAN & KESIMPULAN HASIL PSIKOTES", size=Pt(12))

    # Header + 8 Rows + Final Rec Row
    table_sum = doc.add_table(rows=len(SUMMARY_CONFIGS)+2, cols=6)
    table_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sum.style = 'Table Grid'
    table_sum.autofit = False
    
    # Exact widths scaled to fit A4 (18.5cm total)
    # Hasil kept at 5.0cm as requested previously
    sum_widths = [Cm(0.8), Cm(2.8), Cm(5.0), Cm(5.9), Cm(2.0), Cm(2.0)]
    for r_idx, row in enumerate(table_sum.rows):
        if r_idx == 0: row.height = Cm(0.6)
        for i, width in enumerate(sum_widths):
            row.cells[i].width = width
    
    headers = ["No", "Tes", "Hasil", "Interpretasi", "Status", "Rekomendasi"]
    for i, h in enumerate(headers):
        cell = table_sum.cell(0, i)
        set_cell_background(cell, "156082")
        set_cell_text(cell, h, bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(7.5))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    for i, (code, name) in enumerate(SUMMARY_CONFIGS):
        row = table_sum.rows[i+1]
        row.height = Cm(0.7)
        
        # Col 0: No (Center)
        set_cell_text(row.cells[0], str(i+1), size=Pt(7.5), align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 1: Tes (Left)
        set_cell_text(row.cells[1], name, bold=True, color=COLOR_BLUE_TEXT, size=Pt(7.5), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
        
        det = res_dict.get(code, {}).details if code in res_dict else None
        hasil, interp, status = get_summary_data(code, det)
        
        # Check for manual overrides from the new decision feature
        manual_decisions = (user.report_decisions or {}).get("test_decisions", {}).get(code, {})
        manual_status = manual_decisions.get("status")
        manual_rec = manual_decisions.get("recommendation")

        # Col 2: Hasil (Left)
        set_cell_text(row.cells[2], hasil, size=Pt(7.5), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
        
        # Col 3: Interpretasi (Left)
        set_cell_text(row.cells[3], interp, size=Pt(7.5), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
        row.cells[3].paragraphs[0].runs[0].font.italic = True
        
        # Col 4: Status (Center)
        final_status = manual_status if manual_status else status
        stat_color = RGBColor(16,185,129) # Green
        if final_status in ["UNGGUL", "AMAN", "Lulus", "LANJUT"]: stat_color = COLOR_BLUE_TEXT
        if final_status in ["Tidak Lulus", "TIDAK SESUAI", "TIDAK LANJUT"]: stat_color = RGBColor(220, 38, 38) # Red
        set_cell_text(row.cells[4], final_status, bold=True, color=stat_color, size=Pt(7.5), align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 5: Rekomendasi (Center)
        final_rec = manual_rec if manual_rec else "✔ LANJUT"
        rec_color = RGBColor(6,95,70) # Dark Green
        if final_rec in ["LULUS"]: rec_color = COLOR_BLUE_TEXT
        if final_rec in ["Tidak Lanjut", "TIDAK LANJUT", "TIDAK LULUS"]: rec_color = RGBColor(220, 38, 38) # Red
        set_cell_text(row.cells[5], final_rec, bold=True, color=rec_color, size=Pt(7.5), align=WD_ALIGN_PARAGRAPH.CENTER)
        
        for cell in row.cells: cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Final Recommendation Row
    row_final = table_sum.rows[-1]
    row_final.height = Cm(0.8)
    
    # Merge cells for "Final Recommendation :" label
    c_label = row_final.cells[0]
    c_label.merge(row_final.cells[3])
    set_cell_background(c_label, "F1F5F9")
    set_cell_text(c_label, "Final Recommendation :", bold=True, color=COLOR_BLUE_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Get the saved final status or default to calculated
    # Note: calculateSystemSuggestion logic is on frontend, but we pull the saved 'final_status'
    saved_final = (user.report_decisions or {}).get("final_status", "DIPERTIMBANGKAN")
    
    # "Direkomendasikan" cell (Status, Rekomendasi)
    c_rec = row_final.cells[4]
    c_rec.merge(row_final.cells[5])
    
    # Formatting based on status
    bg_color = "D1FAE5" # Light Green
    text_color = RGBColor(6,95,70) # Dark Green
    display_text = saved_final.title() # e.g. "Direkomendasikan"
    
    if saved_final == "DIPERTIMBANGKAN":
        bg_color = "FEF3C7" # Light Yellow
        text_color = RGBColor(146, 64, 14) # Dark Yellow
    elif saved_final == "TIDAK DISARANKAN":
        bg_color = "FEE2E2" # Light Red
        text_color = RGBColor(153, 27, 27) # Dark Red
        
    set_cell_background(c_rec, bg_color)
    set_cell_text(c_rec, display_text, bold=True, color=text_color, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    for cell in row_final.cells: cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

    # --- 10. KEUNGGULAN & CATATAN PENGEMBANGAN ---
    add_section_heading(doc, "KEUNGGULAN & CATATAN PENGEMBANGAN")
    table_sw = doc.add_table(rows=2, cols=2)
    table_sw.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sw.style = 'Table Grid'
    table_sw.autofit = False
    for row in table_sw.rows:
        row.cells[0].width = Cm(9.24)
        row.cells[1].width = Cm(9.24)
    
    set_cell_background(table_sw.cell(0,0), "D1FAE5") # Light green
    set_cell_text(table_sw.cell(0,0), "KEUNGGULAN KANDIDAT", bold=True, color=RGBColor(6,95,70), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
    
    set_cell_background(table_sw.cell(0,1), "FEF3C7") # Light yellow
    set_cell_text(table_sw.cell(0,1), "YANG PERLU DIPERHATIKAN", bold=True, color=RGBColor(146,64,14), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Set header height and vertical alignment
    table_sw.rows[0].height = Cm(0.6)
    table_sw.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_sw.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Simple placeholders
    k_text = "+\n+\n+"
    p_text = "-\n-\n-"
    
    set_cell_text(table_sw.cell(1,0), k_text, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table_sw.cell(1,1), p_text, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT)
    
    table_sw.cell(1,0).vertical_alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_sw.cell(1,1).vertical_alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # --- 11. SIGNATURE SECTION ---
    table_sig = doc.add_table(rows=4, cols=3)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sig.style = 'Table Grid'
    table_sig.autofit = False
    
    # Row 1: Merged (Validated by) - DARK BLUE HEADER
    r1_cell = table_sig.cell(0, 0)
    r1_cell.merge(table_sig.cell(0, 2))
    set_cell_background(r1_cell, "156082")
    set_cell_text(r1_cell, "Validated by", bold=True, color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
    r1_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_sig.rows[0].height = Cm(0.6)

    # Row 2: Location/Date (Merged)
    r2_cell = table_sig.cell(1, 0)
    r2_cell.merge(table_sig.cell(1, 2))
    set_cell_text(r2_cell, "Location:\nDate:", size=Pt(8), color=RGBColor(100,100,100), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
    r2_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_sig.rows[1].height = Cm(0.8)

    # Row 3: Titles
    table_sig.cell(2, 0).width = Cm(6.17)
    r3_merged = table_sig.cell(2, 1).merge(table_sig.cell(2, 2))
    r3_merged.width = Cm(12.33)
    set_cell_text(table_sig.cell(2, 0), "Diperiksa Oleh:", size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)
    set_cell_text(r3_merged, "Diketahui:", size=Pt(8), align=WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.left_indent = Pt(5)

    # Row 4: Signature Lines
    sig_titles = ["Assessor", "HRGA & CSR Manager", "HRGA Department"]
    for i in range(3):
        cell = table_sig.cell(3, i)
        cell.width = Cm(6.17)
        p = cell.add_paragraph("\n\n\n________________________\n")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sig_titles[i])
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_BLUE_TEXT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_footer.add_run("_DOKUMEN INI BERSIFAT RAHASIA – HANYA UNTUK PENGGUNAAN INTERNAL HRGA DEPARTMENT_")
    r_foot.font.color.rgb = RGBColor(220,38,38)
    r_foot.font.size = Pt(8)
    r_foot.bold = True

    # Return as BytesIO
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()
